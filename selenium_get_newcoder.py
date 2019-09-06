import time
import sys
from selenium import webdriver
from urllib.parse import quote
from selenium.common.exceptions import TimeoutException
from selenium.webdriver.common.by import By 
from selenium.webdriver.support import expected_conditions as EC 
from selenium.webdriver.support.wait import WebDriverWait
from pyquery import PyQuery as pq 
from docx import Document
from docx.shared import Inches
browser = webdriver.Chrome()

base_url = "https://www.nowcoder.com/ta/review-test/review?page=1"

class NewCoder:
    def __init__(self, url):
        browser.get(url)
        browser.maximize_window()

    def get_content_by_page_num(self, pages):
        question_css_selector = "body > div.nk-container > div.final-main.clearfix > div.final-box.clearfix > div:nth-child(1) > div.final-question"
        anwser_css_selector = "body > div.nk-container > div.final-main.clearfix > div.final-box.clearfix > div:nth-child(2) > div.design-answer-box"
        # 创建文档对象
        document = Document()

        for i in range(1, pages+1):
            time.sleep(1)
            input_num = browser.find_element_by_id("turnTo")
            input_num.send_keys(str(i))
            print("-------------------------正在进行第 {} 页获取-------------------".format(i))
            # 点击确定
            time.sleep(1)
            browser.find_element_by_css_selector("body > div.nk-container > div.final-main.clearfix > div.pagination.final-pagination > ul > li.txt-pager > div > span > button").click()
            time.sleep(1)
            # 获取内容
            question = browser.find_element_by_css_selector(question_css_selector).text
            answser = browser.find_element_by_css_selector(anwser_css_selector).text
            time.sleep(0.6)
            print(question)
            # 添加一级标题
            document.add_heading(question, level=1)
            print(answser)
            document.add_paragraph(answser)

        # 保存文档
        document.save('demo1.docx')


ne = NewCoder(base_url)
ne.get_content_by_page_num(413)





#
# def main():
#
#
#     # 创建文档对象
#     document = Document()
#
#     # 设置文档标题，中文要用unicode字符串
#     document.add_heading(u'我的一个新文档', 0)
#
#     # 往文档中添加段落
#     p = document.add_paragraph('This is a paragraph having some ')
#     p.add_run('bold ').bold = True
#     p.add_run('and some ')
#     p.add_run('italic.').italic = True
#
#     # 添加一级标题
#     document.add_heading(u'一级标题, level = 1', level=1)
#     document.add_paragraph('Intense quote', style='IntenseQuote')
#
#     # 添加无序列表
#     document.add_paragraph('first item in unordered list', style='ListBullet')
#
#     # 添加有序列表
#     document.add_paragraph('first item in ordered list', style='ListNumber')
#     document.add_paragraph('second item in ordered list', style='ListNumber')
#     document.add_paragraph('third item in ordered list', style='ListNumber')
#
#     # 添加图片，并指定宽度
#     document.add_picture('e:/docs/pic.png', width=Inches(1.25))
#
#     # 添加表格: 1行3列
#     table = document.add_table(rows=1, cols=3)
#     # 获取第一行的单元格列表对象
#     hdr_cells = table.rows[0].cells
#     # 为每一个单元格赋值
#     # 注：值都要为字符串类型
#     hdr_cells[0].text = 'Name'
#     hdr_cells[1].text = 'Age'
#     hdr_cells[2].text = 'Tel'
#     # 为表格添加一行
#     new_cells = table.add_row().cells
#     new_cells[0].text = 'Tom'
#     new_cells[1].text = '19'
#     new_cells[2].text = '12345678'
#
#     # 添加分页符
#     document.add_page_break()
#
#     # 往新的一页中添加段落
#     p = document.add_paragraph('This is a paragraph in new page.')
#
#     # 保存文档
#     document.save('e:/docs/demo1.docx')
        
        
    










